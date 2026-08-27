"""Turn an uploaded file into plain text.

The bytes are never stored — this is the only thing that ever sees them. A file
whose text cannot be recovered is rejected at upload rather than saved as an
empty row, because a silently-empty knowledge file looks configured but does
nothing, which is the worst outcome for the user.

Blocking, CPU-bound parsing. Call it from a threadpool (`run_in_threadpool`).
"""

import io
import re

from core.logging import get_logger

log = get_logger(__name__)

# Gmail and the LLM are the expensive parts of drafting; a huge upload only
# costs parse time, so the ceiling here is about protecting the request, not the
# storage. 10 MB comfortably covers a long PDF handbook.
MAX_UPLOAD_BYTES = 10 * 1024 * 1024

EXT_PDF = ".pdf"
EXT_DOCX = ".docx"
PLAIN_EXTS = (".txt", ".md", ".markdown", ".text", ".csv")

# Matched on the filename extension rather than the browser-supplied
# content_type: browsers disagree about .md (text/markdown vs application/octet-stream
# vs empty) and .docx, so the extension is the more reliable signal.
SUPPORTED_EXTS = (EXT_PDF, EXT_DOCX, *PLAIN_EXTS)


class ExtractionError(Exception):
    """Base for upload problems the user can act on."""


class UnsupportedFile(ExtractionError):
    pass


class FileTooLarge(ExtractionError):
    pass


class NoTextFound(ExtractionError):
    pass


def _extension(filename: str) -> str:
    _, dot, ext = filename.rpartition(".")
    return f".{ext.lower()}" if dot else ""


def _tidy(text: str) -> str:
    """Collapse the whitespace PDF extraction leaves behind.

    Extractors emit a newline per layout line, so a paragraph arrives as a dozen
    short lines and page breaks as runs of blank ones. Left alone this wastes a
    large share of the prompt budget on whitespace.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for index, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            # One malformed page should not lose the other ninety-nine.
            log.warning("drafts.pdf_page_failed", page=index)
    return "\n\n".join(p for p in pages if p.strip())


def _from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables carry real content in the documents people upload here — price
    # lists, contact tables, FAQ grids — so they are not skipped.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_plain(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        # A file exported from Windows tooling is usually cp1252; latin-1 never
        # raises, so this is the terminal fallback.
        return data.decode("latin-1", errors="replace")


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an upload. Raises `ExtractionError` on bad input."""
    if len(data) > MAX_UPLOAD_BYTES:
        raise FileTooLarge(
            f"file is {len(data) // 1_048_576} MB; the limit is {MAX_UPLOAD_BYTES // 1_048_576} MB"
        )
    if not data:
        raise NoTextFound("the file is empty")

    ext = _extension(filename)
    if ext == EXT_PDF:
        raw = _from_pdf(data)
    elif ext == EXT_DOCX:
        raw = _from_docx(data)
    elif ext in PLAIN_EXTS:
        raw = _from_plain(data)
    else:
        raise UnsupportedFile(
            f"{filename!r} is not a supported file type; upload one of {', '.join(SUPPORTED_EXTS)}"
        )

    text = _tidy(raw)
    if not text:
        # Overwhelmingly a scanned PDF: pages are images, so there is no text
        # layer to read. Say so, because "upload failed" would send the user
        # looking for the wrong problem.
        raise NoTextFound(
            f"no text could be read from {filename!r}. If it is a scan or a "
            "photo, it has no text layer — run OCR on it first, or paste the "
            "text into custom instructions instead."
        )
    return text
